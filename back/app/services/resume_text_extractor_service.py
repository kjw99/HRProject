import re
import subprocess
import tempfile
from pathlib import Path

import fitz
from docx import Document


class ResumeTextExtractorService:
    def extract_text(self, path: str | Path) -> str:
        file_path = Path(path)
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            return self._extract_pdf(file_path)

        if suffix == ".docx":
            return self._extract_docx(file_path)

        if suffix == ".hwp":
            return self._extract_hwp(file_path)

        if suffix in {".txt", ".md"}:
            return self._extract_plain_text(file_path)

        if suffix in {".doc", ".ppt", ".pptx"}:
            return self._extract_via_libreoffice(file_path)

        raise ValueError(
            f"지원하지 않는 이력서 파일 형식입니다: {suffix or '알 수 없음'}"
        )

    def _extract_pdf(self, path: Path) -> str:
        parts: list[str] = []
        with fitz.open(path) as document:
            for page in document:
                page_text = page.get_text("text")
                if page_text:
                    parts.append(page_text)

        return self._normalize_text("\n".join(parts))

    def _extract_docx(self, path: Path) -> str:
        document = Document(path)
        parts: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return self._normalize_text("\n".join(parts))

    def _extract_hwp(self, path: Path) -> str:
        try:
            from hwp5.filestructure import Hwp5File
            from hwp5.recordstream import read_records
        except ImportError as exc:
            raise ValueError("HWP 텍스트 추출을 위해 pyhwp 패키지가 필요합니다.") from exc

        document = Hwp5File(str(path))
        parts: list[str] = []

        for section_name in document.text:
            section_stream = document.text[section_name].open()
            for record in read_records(section_stream):
                if record.get("tagid") != 67:
                    continue

                payload = record.get("payload", b"")
                if not payload:
                    continue

                decoded = payload.decode("utf-16-le", errors="ignore")
                cleaned = "".join(
                    character if character >= " " or character in "\n\r\t" else " "
                    for character in decoded
                )
                cleaned = cleaned.strip()
                if cleaned:
                    parts.append(cleaned)

        return self._normalize_text("\n".join(parts))

    def _extract_plain_text(self, path: Path) -> str:
        for encoding in ("utf-8-sig", "utf-8", "cp949"):
            try:
                return self._normalize_text(path.read_text(encoding=encoding))
            except UnicodeDecodeError:
                continue

        return self._normalize_text(path.read_text(errors="ignore"))

    def _extract_via_libreoffice(self, path: Path) -> str:
        with tempfile.TemporaryDirectory() as temp_dir:
            result = subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    temp_dir,
                    str(path),
                ],
                check=False,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
            if result.returncode != 0:
                raise ValueError(
                    "LibreOffice로 이력서를 PDF로 변환하지 못했습니다."
                )

            converted_path = Path(temp_dir) / f"{path.stem}.pdf"
            if not converted_path.exists():
                matches = list(Path(temp_dir).glob("*.pdf"))
                if not matches:
                    raise ValueError("변환된 PDF 파일이 생성되지 않았습니다.")

                converted_path = matches[0]

            return self._extract_pdf(converted_path)

    def _normalize_text(self, value: str) -> str:
        normalized = value.replace("\r\n", "\n").replace("\r", "\n")
        normalized = re.sub(r"[ \t]+", " ", normalized)
        normalized = re.sub(r"\n{3,}", "\n\n", normalized)
        return normalized.strip()


resume_text_extractor_service = ResumeTextExtractorService()
